from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Security, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security.api_key import APIKeyHeader
from pydantic import BaseModel, Field
import asyncio
from typing import List, Optional, Dict, Any
import os
from dotenv import load_dotenv

load_dotenv(os.path.join(os.path.dirname(__file__), "..", ".env"))

# --- MAC OS MULTIPROCESSING FIX ---
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["OMP_NUM_THREADS"] = "1"

import uvicorn
import shutil
import numpy as np
import time
import uuid
import json
import logging
from collections import defaultdict
import platform
import datetime

start_time = time.time()

# Identify if Mac ARM for hardware acceleration
mac_device = "mps" if platform.system() == "Darwin" and platform.machine() == "arm64" else "cpu"

from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
import faiss
from pymilvus import MilvusClient
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.messages import HumanMessage, AIMessage
from rank_bm25 import BM25Okapi
from fastapi.concurrency import run_in_threadpool
from core_agent import create_agent
from outcome_predictor import OutcomePredictor
from performance_optimizer import perf_monitor
from llm_utils import get_llm
import core_agent
import outcome_predictor
import enhanced_feature_extractor
import enhanced_training_data
import performance_optimizer
import functools

def timing_decorator(func):
    """Decorator that records endpoint call timing using perf_monitor."""
    @functools.wraps(func)
    async def wrapper(*args, **kwargs):
        perf_monitor.start_timer()
        try:
            result = await func(*args, **kwargs)
            perf_monitor.end_timer("api_calls")
            return result
        except Exception:
            perf_monitor.end_timer("api_calls")
            raise
    return wrapper

# --- 1. GOVERNANCE & LOGGING ---
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] [LuminaAPI] %(message)s",
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger("lumina_api")

app = FastAPI(title="Legal AI Research Engine API", version="2.3")

# --- 2. SECURITY (CORS) ---
ALLOWED_ORIGINS = os.environ.get("ALLOWED_ORIGINS", "http://localhost:5173,http://localhost:3000").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- 3. SECURITY (API KEY AUTH) ---
API_KEY = os.environ.get("LUMINA_API_KEY")
api_key_header = APIKeyHeader(name="X-API-Key", auto_error=False)

from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
security = HTTPBearer(auto_error=False)

async def verify_api_key(
    api_key: str = Security(api_key_header),
    auth: Optional[HTTPAuthorizationCredentials] = Depends(security)
):
    # Use key from X-API-Key or Authorization: Bearer
    token = api_key
    if auth:
        token = auth.credentials

    if not API_KEY:
        # In development, if no key is set, we might allow but warn
        logger.warning("LUMINA_API_KEY not set in environment. Running in unsecured mode.")
        return "development"
    
    if token and token != API_KEY:
        logger.warning("Attempted access with invalid API key.")
        raise HTTPException(status_code=403, detail="Invalid API Key")
    elif not token:
        logger.warning("Access without API key. Rejecting.")
        raise HTTPException(status_code=401, detail="API Key missing")
    return token

# --- 4. SECURITY (RATE LIMITING) ---
class SimpleRateLimiter:
    def __init__(self, limit=20, window=60):
        self.requests = defaultdict(list)
        self.limit = limit
        self.window = window

    def __call__(self, request: Request, api_key: str = Depends(verify_api_key)):
        now = time.time()
        client_id = api_key if api_key != "anonymous" else request.client.host
        self.requests[client_id] = [req for req in self.requests[client_id] if now - req < self.window]
        if len(self.requests[client_id]) >= self.limit:
            logger.warning(f"Rate limit exceeded for client {client_id}")
            raise HTTPException(status_code=429, detail="Rate limit exceeded. Try again later.")
        self.requests[client_id].append(now)
        return client_id

# Global State
class AppState:
    vector_store = None
    embeddings_model = None
    bm25_index = None
    bm25_corpus = None
    semantic_cache_store = None # FAISS-backed semantic cache

state = AppState()

# Models
class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1, max_length=4000)
    history: List[Dict[str, str]] = []
    language: str = "en"
    jurisdiction: str = "All"

class ChatResponse(BaseModel):
    response: str
    sources: List[Dict[str, str]] = []
    cached: bool = False

class PredictionRequest(BaseModel):
    description: str
    court: Optional[str] = None
    judge: Optional[str] = None
    case_type: Optional[str] = None
    model_version: str = "advanced"

class PredictionResponse(BaseModel):
    result: Dict[str, Any]

class MilvusLiteWrapper:
    """Wrapper for MilvusClient to mimic LangChain VectorStore interface for similarity_search"""
    def __init__(self, uri, collection_name, embedding_func):
        self.client = MilvusClient(uri=uri)
        self.collection_name = collection_name
        self.embedding_func = embedding_func

    def similarity_search(self, query: str, k: int = 4, **kwargs) -> List[Any]:
        # Generate embedding
        query_vec = self.embedding_func.embed_query(query)
        
        filter_expr = kwargs.get('filter')
        
        # Search
        res = self.client.search(
            collection_name=self.collection_name,
            data=[query_vec],
            limit=k,
            filter=filter_expr,
            output_fields=["text_content", "filename", "page_number", "modality", "jurisdiction"]
        )
        
        documents = []
        if not res:
            return documents
            
        for hit in res[0]:
            entity = hit['entity']
            # Milvus Lite returns entity in 'entity' key usually, or at top level depending on version?
            # MilvusClient search result is list of list of dicts.
            # Dict keys: id, distance, entity={...}
            
            content = entity.get('text_content', '')
            meta = {
                "source": entity.get('filename'),
                "page": entity.get('page_number'),
                "modality": entity.get('modality'),
                "score": hit.get('distance')
            }
            documents.append(Document(page_content=content, metadata=meta))
            
        return documents

    @property
    def docstore(self):
        # Mocking docstore for BM25 (simplified)
        # We can't easily fetch all docs from Milvus without iterator. 
        # Returning empty to skip BM25 for now or implement scroll if needed.
        class MockDocstore:
            _dict = {}
        return MockDocstore()

def update_bm25_index():
    """Helper to rebuild BM25 index from Vector Store documents"""
    if state.vector_store and hasattr(state.vector_store, 'docstore'):
        try:
            print("🔄 Building BM25 Index...")
            # Extract all docs
            docs = list(state.vector_store.docstore._dict.values())
            if not docs:
                return
            
            # Tokenize
            tokenized_corpus = [doc.page_content.lower().split() for doc in docs]
            
            # Build Index
            state.bm25_index = BM25Okapi(tokenized_corpus)
            state.bm25_corpus = docs # Store actual doc objects for retrieval
            print(f"✅ BM25 Index Built ({len(docs)} documents)")
        except Exception as e:
            print(f"⚠️ Failed to build BM25 Index: {e}")

# Startup
@app.on_event("startup")
async def startup_event():
    print("🚀 Starting Legal AI Engine...")
    
    # Startup Checks
    if not os.environ.get("INDIAN_KANOON_API_TOKEN"):
        logger.warning("INDIAN_KANOON_API_TOKEN is missing. Indian Kanoon search will be disabled. (Optional)")
        
    try:
        # Load Embeddings
        try:
            # Check if InLegalBERT feature extractor initialized it
            # We use langchain wrapper here for vector store compatibility
            state.embeddings_model = HuggingFaceEmbeddings(
                model_name="law-ai/InLegalBERT", # Match the extractor
                model_kwargs={"device": mac_device}
            )
            print(f"✅ Embeddings Model Loaded (InLegalBERT on {mac_device})")
        except Exception:
            print("⚠️ InLegalBERT loading failed. Falling back to all-MiniLM-L6-v2.")
            state.embeddings_model = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={"device": mac_device}
            )
        
        # Load Vector Store
        milvus_db_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "services", "ingestion", "milvus_demo.db"))
        print(f"DEBUG: Checking for Milvus DB at {milvus_db_path}")
        if os.path.exists(milvus_db_path):
             print(f"✅ Found Milvus DB at {milvus_db_path}")
             # Use the same embedding model as ingestion
             state.embeddings_model = HuggingFaceEmbeddings(
                model_name="sentence-transformers/clip-ViT-B-32-multilingual-v1",
                model_kwargs={"device": mac_device}
             )
             print(f"✅ Embeddings Model Loaded (CLIP Multilingual on {mac_device})")
             
             try:
                 state.vector_store = MilvusLiteWrapper(
                    uri=milvus_db_path,
                    collection_name="legal_rag_multimodal",
                    embedding_func=state.embeddings_model
                 )
                 print("✅ Milvus Vector Store Loaded (Custom Wrapper)")
             except Exception as e:
                 print(f"❌ Failed to load Milvus Wrapper: {e}")
                 
        if state.vector_store is None and os.path.exists("faiss_store"):
            # Fallback to legacy FAISS
            # Load default embeddings for FAISS if not CLIP
             state.embeddings_model = HuggingFaceEmbeddings(
                model_name="sentence-transformers/clip-ViT-B-32-multilingual-v1", 
                model_kwargs={"device": "cpu"}
            )
             state.vector_store = FAISS.load_local(
                "faiss_store", state.embeddings_model, allow_dangerous_deserialization=True
            )
             print("✅ Vector Store Loaded (FAISS)")
            
            # Build BM25
             update_bm25_index()
            
        if not state.vector_store:
            logger.warning("Vector Store not found. Upload PDFs to initialize.")
            
        # --- SEMANTIC CACHE INIT ---
        cache_path = "semantic_cache_index"
        if os.path.exists(cache_path):
            state.semantic_cache_store = FAISS.load_local(
                cache_path, state.embeddings_model, allow_dangerous_deserialization=True
            )
            logger.info(f"✅ Loaded FAISS Semantic Cache from {cache_path}")
        else:
            # Initialize empty FAISS
            empty_doc = Document(page_content="[cache_init]", metadata={"response": "{}"})
            state.semantic_cache_store = FAISS.from_documents([empty_doc], state.embeddings_model)
            logger.info("✅ Initialized new FAISS Semantic Cache")
            
        # --- MODEL PRE-LOADING ---
        from core_agent import get_predictor
        try:
            get_predictor()
            logger.info("✅ ML Predictive Model Pre-loaded")
        except Exception as me:
            logger.error(f"❌ Failed to pre-load ML Model: {me}")

    except Exception as e:
        logger.error(f"Error loading resources: {e}")

# Endpoints
@app.get("/health")
async def health():
    try:
        from core_agent import ACTIVE_LLM_PROVIDER, ACTIVE_LLM_MODEL, GLOBAL_MODELS
        llm_string = f"{ACTIVE_LLM_PROVIDER}/{ACTIVE_LLM_MODEL}" if ACTIVE_LLM_PROVIDER != "unknown" else "ollama/llama3.2"
        return {
            "status": "ok",
            "vector_db": "milvus_lite" if state.vector_store is not None else "unavailable",
            "bm25_loaded": state.bm25_index is not None,
            "ml_model_loaded": GLOBAL_MODELS.predictor is not None,
            "docs_indexed": len(state.bm25_corpus) if state.bm25_corpus else 0,
            "llm": llm_string
        }
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/llm-status")
async def llm_status():
    from core_agent import ACTIVE_LLM_PROVIDER, ACTIVE_LLM_MODEL
    return {"provider": ACTIVE_LLM_PROVIDER, "model": ACTIVE_LLM_MODEL}

@app.post("/chat", response_model=ChatResponse)
@timing_decorator
async def chat_endpoint(request: ChatRequest, client_id: str = Depends(SimpleRateLimiter(limit=30, window=60))):
    try:
        # --- TRANSLATION (Source to English) ---
        if request.language and request.language != 'en':
            try:
                from deep_translator import GoogleTranslator
                translated = GoogleTranslator(source='auto', target='en').translate(request.message)
                request.message = translated
                logger.info(f"Translated to English: {request.message}")
            except Exception as e:
                logger.warning(f"Translation to English failed: {e}")

        # --- JURISDICTION CONTEXT ---
        if request.jurisdiction and request.jurisdiction != "All":
            request.message = f"[Jurisdiction Context: {request.jurisdiction}] {request.message}"
            
        # --- SEMANTIC CACHE LOOKUP (FAISS) ---
        if state.semantic_cache_store:
            # We use a tight L2 threshold for "very similar" matches against the persistent FAISS index
            results = state.semantic_cache_store.similarity_search_with_score(request.message, k=1)
            if results:
                doc, l2_distance = results[0]
                if doc.page_content != "[cache_init]" and l2_distance < 0.25:
                    logger.info(f"⚡ Cache Hit (L2 Distance: {l2_distance:.4f})")
                    try:
                        cached_data = json.loads(doc.metadata["response"])
                        return ChatResponse(
                            response=cached_data["response"], 
                            sources=cached_data.get("sources", []), 
                            cached=True
                        )
                    except Exception as e:
                        logger.warning(f"Failed to parse cached response: {e}")
        # --- END CACHE ---

        # Initialize Agent
        agent_executor = create_agent(
            state.vector_store,
            bm25_index=state.bm25_index,
            bm25_corpus=state.bm25_corpus
        )
        
        # Convert history
        langchain_history = []
        for msg in request.history[-4:]:  # Limit to last 4 to avoid translation latency
            content = msg['content']
            if request.language and request.language != 'en':
                try:
                    from deep_translator import GoogleTranslator
                    content = GoogleTranslator(source='auto', target='en').translate(content)
                except Exception as e:
                    print(f"⚠️ History translation failed: {e}")
            
            if msg['role'] == 'user':
                langchain_history.append(HumanMessage(content=content))
            elif msg['role'] == 'assistant':
                langchain_history.append(AIMessage(content=content))
        
        langchain_history.append(HumanMessage(content=request.message))
        
        initial_state = {"messages": langchain_history, "step_count": 0}
        
        try:
            final_state = await asyncio.wait_for(
                run_in_threadpool(agent_executor.invoke, initial_state), 
                timeout=60.0
            )
        except asyncio.TimeoutError:
            raise HTTPException(status_code=504, detail="The legal AI agent took too long to respond. The request has timed out.")
        
        final_message = final_state["messages"][-1].content
        
        # --- TRANSLATION (English to Target) ---
        if request.language and request.language != 'en':
            try:
                from deep_translator import GoogleTranslator
                # deep-translator handles limits up to 5000 chars natively for Google Translate
                final_message = GoogleTranslator(source='en', target=request.language).translate(final_message)
                print(f"🌐 Translated response back to target language")
            except Exception as e:
                print(f"⚠️ Translation to {request.language} failed: {e}")
        
        sources = []
        for msg in final_state["messages"]:
            if hasattr(msg, "tool_calls") and msg.tool_calls:
                for tool_call in msg.tool_calls:
                    sources.append({
                        "type": tool_call['name'],
                        "args": str(tool_call['args'])
                    })
        
        response_obj = ChatResponse(response=final_message, sources=sources, cached=False)
        
        # --- UPDATE CACHE (FAISS) ---
        if state.semantic_cache_store:
            cache_meta = {
                "response": json.dumps({
                    "response": final_message,
                    "sources": sources
                })
            }
            cache_doc = Document(page_content=request.message, metadata=cache_meta)
            state.semantic_cache_store.add_documents([cache_doc])
            
            try:
                state.semantic_cache_store.save_local("semantic_cache_index")
            except Exception as e:
                logger.error(f"Failed to save semantic cache: {e}")
        
        return response_obj
        
    except Exception as e:
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

class SimilarityRequest(BaseModel):
    description: str
    jurisdiction: Optional[str] = "All"

@app.post("/similar_cases")
@timing_decorator
async def similar_cases_endpoint(request: SimilarityRequest, client_id: str = Depends(SimpleRateLimiter(limit=30, window=60))):
    try:
        if not state.vector_store:
            raise HTTPException(status_code=500, detail="Vector store not initialized. Upload documents first.")
            
        # We need raw similarity search to get the Document objects.
        # Handle the syntactic difference between Langchain Milvus (expr="string") and FAISS (filter={"dict"})
        is_faiss = hasattr(state.vector_store, 'index') or state.vector_store.__class__.__name__ == 'FAISS'
        
        if request.jurisdiction and request.jurisdiction != "All":
            if is_faiss:
                similar_docs = state.vector_store.similarity_search(request.description, k=5, filter={"jurisdiction": request.jurisdiction})
            else:
                similar_docs = state.vector_store.similarity_search(request.description, k=5, filter=f"jurisdiction == '{request.jurisdiction}'")
        else:
            similar_docs = state.vector_store.similarity_search(request.description, k=5)
        
        if not similar_docs:
            return {
                "analytics": {
                    "winRate": 0,
                    "avgDuration": "N/A",
                    "judgeTendency": "Unknown",
                    "outcomes": [
                        {"name": "Allowance", "value": 0, "color": "#4F46E5"},
                        {"name": "Dismissal", "value": 0, "color": "#E11D48"},
                        {"name": "Settlement", "value": 0, "color": "#94A3B8"}
                    ]
                },
                "cases": []
            }
            
        # 2. Extract snippets and metadata for the frontend
        precedents = []
        context_blocks = []
        for i, doc in enumerate(similar_docs):
            source = doc.metadata.get('source', 'Unknown Document')
            page = doc.metadata.get('page', '?')
            snippet = doc.page_content[:800] # Limit size
            
            precedents.append({
                "id": i + 1,
                "source": source,
                "page": page,
                "snippet": snippet
            })
            
            context_blocks.append(f"CASE {i+1} [{source}]:\n{snippet}\n")
            
        # 3. Use ChatGroq to analyze the outcomes of these specific cases and generate structured UI data
        # 3. Use ChatGroq to analyze the outcomes of these specific cases and generate structured UI data
        from langchain_groq import ChatGroq
        from langchain_core.prompts import ChatPromptTemplate
        from langchain_core.output_parsers import JsonOutputParser
        from pydantic import BaseModel, Field
        from typing import List
        import os
        
        class PrecedentCase(BaseModel):
            id: int
            name: str
            year: int
            court: str
            match: int
            outcome: str
            duration_months: int
            factSimilarity: str
            legalSimilarity: str
            tags: List[str]
            reason: str
            
        class CasesList(BaseModel):
            cases: List[PrecedentCase]

        parser = JsonOutputParser(pydantic_object=CasesList)

        llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0.1)
        
        analysis_prompt = """You are a legal expert analyzing historical precedents for a Prediction Dashboard.
Given the user's case description and {num_cases} similar historical cases retrieved from our database, generate a structured JSON object containing a list of 'cases'.

Extract and synthesize ONLY the 'cases' array for the specific {num_cases} precedents provided below.
- id: integer starting from 1
- name: deduce a case name from the snippet (e.g. "State v. Sharma")
- year: integer (guess based on context or use recent year)
- court: court name
- match: semantic match percentage (0-100)
- outcome: EXACTLY ONE OF "Allowance", "Dismissal", or "Settlement"
- duration_months: integer, estimate resolution time in months based on dates in context, or 0 if unknown
- factSimilarity: "High", "Medium", or "Low"
- legalSimilarity: "High", "Medium", or "Low"
- tags: Array of 2-3 legal tags (e.g. ["Bail", "PMLA"])
- reason: 1-2 sentences explaining the court's reasoning and ratio decidendi based on the snippet.

User's Case: {user_desc}

--- RETRIEVED PRECEDENTS ---
{context}

CRITICAL: Return ONLY valid JSON format matching the schema rules.
{format_instructions}
"""
        prompt = ChatPromptTemplate.from_template(analysis_prompt)
        chain = prompt | llm | parser
        
        try:
            output_json = chain.invoke({
                "num_cases": len(similar_docs),
                "user_desc": request.description,
                "context": "".join(context_blocks),
                "format_instructions": parser.get_format_instructions()
            })
            
            if isinstance(output_json, dict):
                cases_list = output_json.get("cases", [])
            elif isinstance(output_json, list):
                cases_list = output_json
            else:
                cases_list = []
                
            total_cases = len(cases_list)
            
            allowance_count = 0
            dismissal_count = 0
            settlement_count = 0
            total_duration = 0
            cases_with_duration = 0
            
            for c in cases_list:
                outcome = c.get("outcome", "Dismissal")
                if "allowance" in outcome.lower(): 
                    allowance_count += 1
                    c["outcome"] = "Allowance"
                elif "dismissal" in outcome.lower(): 
                    dismissal_count += 1
                    c["outcome"] = "Dismissal"
                else: 
                    settlement_count += 1
                    c["outcome"] = "Settlement"
                
                duration = c.get("duration_months", 0)
                if isinstance(duration, int) and duration > 0:
                    total_duration += duration
                    cases_with_duration += 1
            
            if total_cases > 0:
                win_rate = round((allowance_count / total_cases) * 100)
                dismissal_pct = round((dismissal_count / total_cases) * 100)
                settlement_pct = 100 - win_rate - dismissal_pct
            else:
                win_rate = 0
                dismissal_pct = 0
                settlement_pct = 0
            
            avg_duration_str = f"{int(total_duration / cases_with_duration)} months" if cases_with_duration > 0 else "14 months"
            judge_tendency = "Pro-Allowance" if allowance_count >= dismissal_count else "Strict"
            
            # Use Outcome Predictor to get SHAP and confidence
            explanation = {}
            try:
                from outcome_predictor import OutcomePredictor
                from enhanced_feature_extractor import EnhancedFeatureExtractor
                
                # Assume feature extraction works locally for this snippet
                feat_ex = EnhancedFeatureExtractor()
                features = feat_ex.extract_all_features(request.description)
                features['court'] = request.jurisdiction
                
                # Get predictor from core_agent
                import core_agent
                predictor = core_agent.get_predictor()
                if predictor:
                    pred_res = predictor.predict(features)
                    # Use the new structure
                    explanation = {
                        "predicted_outcome": pred_res.get('predicted_outcome'),
                        "outcome_probability": pred_res.get('outcome_probability'),
                        "top_features": [
                            {"feature": f['feature'], "value": f['impact'], "direction": f['direction']}
                            for f in pred_res.get('top_factors', [])
                        ],
                        "confidence_band": {
                            "low": max(0, pred_res.get('outcome_probability', 0) - 0.1),
                            "high": min(1, pred_res.get('outcome_probability', 0) + 0.1)
                        }
                    }
            except Exception as ml_e:
                logger.warning(f"Failed to generate prediction explanation: {ml_e}")

            # Embed real similar precedent IDs
            if explanation:
                explanation["similar_precedents"] = [c.id for c in cases_list[:3]]
            
            result_json = {
                "analytics": {
                    "winRate": win_rate,
                    "avgDuration": avg_duration_str,
                    "judgeTendency": judge_tendency,
                    "outcomes": [
                        {"name": "Allowance", "value": win_rate, "color": "#4F46E5"},
                        {"name": "Dismissal", "value": dismissal_pct, "color": "#E11D48"},
                        {"name": "Settlement", "value": settlement_pct, "color": "#94A3B8"}
                    ]
                },
                "explanation": explanation,
                "cases": cases_list
            }
            return result_json
        except Exception as e:
            logger.error(f"Failed to parse Groq response: {e}")
            raise HTTPException(status_code=500, detail="Failed to synthesize precedent analytics.")
            
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

PREDICTION_EXPLANATION_PROMPT = """
You are a senior legal strategist. Analyze the following machine learning prediction for a legal case.
The model has predicted the outcome: {predicted_outcome} with {confidence_score_pct}% confidence.

Top Factors influencing the model:
{top_factors}

Case Description:
{description}

Metadata:
{metadata}

TASK:
1. Explain in 3-4 professional sentences WHY the model likely chose this outcome based on the factors and case facts.
2. Highlight any potential risks or "red flags" (e.g. strict judge, unfavorable precedent).
3. Provide a strategic recommendation (e.g. "Focus on distinguishing this from X vs Y precedent").

Return a structured response that is concise and adds strategic value beyond the raw score.
"""

@app.post("/predict", response_model=PredictionResponse)
@timing_decorator
async def predict_outcome(request: PredictionRequest, client_id: str = Depends(SimpleRateLimiter(limit=20, window=60))):
    try:
        from outcome_predictor import OutcomePredictor
        from enhanced_feature_extractor import EnhancedFeatureExtractor
        
        # 1. Extract Features
        feat_ex = EnhancedFeatureExtractor()
        features = feat_ex.extract_all_features(request.description)
        # Override with user provided metadata if available
        if request.court: features['court'] = request.court
        if request.judge: features['judge'] = request.judge
        if request.case_type: features['case_type'] = request.case_type
        
        # 2. Get ML Prediction
        from core_agent import get_predictor
        predictor = get_predictor()
        if not predictor:
            raise HTTPException(status_code=500, detail="Prediction model not initialized.")
            
        ml_res = predictor.predict(features)
        
        # 3. Get Similar Cases for RAG context and precedents
        try:
            if state.vector_store:
                similar_docs = state.vector_store.similarity_search(request.description, k=3)
                cases_list = []
                for i, doc in enumerate(similar_docs):
                    cases_list.append({
                        "id": i + 1,
                        "name": doc.metadata.get('source', 'Unknown Precedent'),
                        "year": 2023, # Placeholder or extracted
                        "court": doc.metadata.get('jurisdiction', 'High Court'),
                        "match": 90 - (i * 5),
                        "outcome": "Allowance" if i == 0 else "Dismissal",
                        "factSimilarity": "High" if i == 0 else "Medium",
                        "legalSimilarity": "High",
                        "tags": ["Precedent", "Commercial"],
                        "reason": doc.page_content[:200]
                    })
                ml_res['cases'] = cases_list
            else:
                ml_res['cases'] = []
        except Exception as e:
            logger.warning(f"Similarity search failed in /predict: {e}")
            ml_res['cases'] = []

        # 4. LLM-Augmented Explanation
        llm, provider = get_llm()
        
        from langchain_core.prompts import ChatPromptTemplate
        prompt = ChatPromptTemplate.from_template(PREDICTION_EXPLANATION_PROMPT)
        
        top_factors_str = "\n".join([f"- {f['feature']}: {f['impact']} ({f['direction']})" for f in ml_res.get('top_factors', [])])
        
        chain = prompt | llm
        llm_explanation_res = await run_in_threadpool(
            chain.invoke, 
            {
                "predicted_outcome": ml_res.get('predicted_outcome', 'unknown'),
                "confidence_score_pct": round(ml_res.get('outcome_probability', 0) * 100, 1),
                "top_factors": top_factors_str,
                "description": request.description,
                "metadata": json.dumps({
                    "court": request.court,
                    "judge": request.judge,
                    "case_type": request.case_type
                })
            }
        )
        
        # 5. Final Response Construction (Step 3B structure)
        final_result = {
            "predicted_outcome": ml_res.get('predicted_outcome', 'unknown'),
            "outcome_probability": ml_res.get('outcome_probability', 0),
            "confidence_score": ml_res.get('confidence_score', 0),
            "top_factors": ml_res.get('top_factors', []),
            "llm_explanation": llm_explanation_res.content,
            "model": provider,
            "cases": ml_res.get('cases', []),
            "analytics": {
                "winRate": round(ml_res.get('outcome_probability', 0) * 100),
                "avgDuration": "12 months",
                "judgeTendency": "Neutral",
                "outcomes": [
                    {"name": 'Allowance', "value": round(ml_res.get('outcome_probability', 0) * 100), "color": '#4F46E5'},
                    {"name": 'Dismissal', "value": 100 - round(ml_res.get('outcome_probability', 0) * 100), "color": '#E11D48'}
                ]
            }
        }
        
        return {"result": final_result}

        
    except Exception as e:
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/upload")
@timing_decorator
async def upload_document(files: List[UploadFile] = File(...), client_id: str = Depends(SimpleRateLimiter(limit=30, window=60))):
    try:
        temp_dir = "temp_uploads"
        os.makedirs(temp_dir, exist_ok=True)
        
        documents = []
        MAX_FILE_SIZE = 20 * 1024 * 1024 # 20 MB
        
        for file in files:
            if not file.filename.lower().endswith('.pdf') or file.content_type != "application/pdf":
                raise HTTPException(status_code=400, detail=f"File {file.filename} fails validation. Only PDF files are permitted.")
                
            content = await file.read()
            if len(content) > MAX_FILE_SIZE:
                raise HTTPException(status_code=413, detail=f"File size exceeds 20MB limit.")
            
            # Check magic bytes for PDF (%PDF)
            if not content.startswith(b'%PDF'):
                raise HTTPException(status_code=400, detail=f"File {file.filename} is corrupted or not a valid PDF document.")
                
            file_path = os.path.join(temp_dir, file.filename)
            with open(file_path, "wb") as f:
                f.write(content)
            
            import fitz  # pymupdf
            doc = fitz.open(stream=content, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            docs = [Document(page_content=text, metadata={"source": file.filename})]
            documents.extend(docs)
            
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        final_documents = text_splitter.split_documents(documents)
        
        is_faiss = hasattr(state.vector_store, 'index') or state.vector_store.__class__.__name__ == 'FAISS'
        if state.vector_store and is_faiss:
            state.vector_store.add_documents(final_documents)
            state.vector_store.save_local("faiss_store")
        else:
            # Fallback to FAISS specifically for uploads if Milvus is default
            emb_model = HuggingFaceEmbeddings(model_name="sentence-transformers/clip-ViT-B-32-multilingual-v1", model_kwargs={"device": "cpu"})
            if os.path.exists("faiss_store"):
                faiss_store = FAISS.load_local("faiss_store", emb_model, allow_dangerous_deserialization=True)
                faiss_store.add_documents(final_documents)
                faiss_store.save_local("faiss_store")
            else:
                faiss_store = FAISS.from_documents(final_documents, emb_model)
                faiss_store.save_local("faiss_store")
        
        # Rebuild BM25 after upload
        update_bm25_index()
        
        # Run LegalNER for auto-tagging on the chunks
        tags = {"acts": set(), "sections": set(), "court": None, "type": "Commercial"}
        try:
            from core_agent import get_ner
            ner = get_ner()
            sample_text = " ".join([d.page_content for d in final_documents[:5]])
            entities = ner.extract_entities(sample_text)
            for e in entities:
                if e['label'] == 'STATUTE':
                    tags['acts'].add(e['text'])
                elif e['label'] == 'PROVISION':
                    tags['sections'].add(e['text'])
                elif e['label'] == 'COURT':
                    tags['court'] = e['text']
            
            # Simple heuristic for case type
            if "arbitration" in sample_text.lower():
                tags['type'] = "Arbitration"
            elif "patent" in sample_text.lower() or "trademark" in sample_text.lower():
                tags['type'] = "Intellectual Property"
            
            tags['acts'] = list(tags['acts'])
            tags['sections'] = list(tags['sections'])
        except Exception as e:
            logger.warning(f"NER failed during upload: {e}")
        
        shutil.rmtree(temp_dir)
        
        return {
            "message": f"Processed {len(files)} files successfully.", 
            "chunks": len(final_documents),
            "tags": {
                "acts": tags['acts'],
                "sections": tags['sections'],
                "court": tags['court'] or "Unknown Court",
                "type": tags['type'],
                "doc_id": str(uuid.uuid4())
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/visualize/timeline")
@timing_decorator
async def visualize_timeline(request: dict):
    """
    Extract chronological timeline from legal case text.
    
    Input: {"text": "..."}
    Output: {"events": [...], "message": "..."}
    """
    try:
        case_text = request.get("text", "")
        
        if not case_text or len(case_text.strip()) < 50:
            raise HTTPException(status_code=400, detail="Please provide sufficient case text (minimum 50 characters)")
        
        from timeline_extractor import TimelineExtractor
        
        extractor = TimelineExtractor()
        # Use run_in_threadpool to prevent the synchronous LangChain/Ollama call from blocking the FastAPI event loop
        timeline_events = await run_in_threadpool(extractor.extract_chronology, case_text)
        
        if not timeline_events:
            return {
                "events": [],
                "message": "No dated events found in the provided text. Please ensure the text contains specific dates or time references."
            }
        
        return {
            "events": timeline_events,
            "message": f"Successfully extracted {len(timeline_events)} timeline events"
        }
        
    except ValueError as e:
        raise HTTPException(status_code=500, detail=f"Configuration error: {str(e)}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Timeline extraction failed: {str(e)}")

class CitationRequest(BaseModel):
    chunk_ids: List[str]

@app.post("/resolve-citations")
async def resolve_citations(request: CitationRequest):
    from core_agent import CITATION_STORE
    results = []
    for cid in request.chunk_ids:
        if cid in CITATION_STORE:
            results.append(CITATION_STORE[cid])
    return {"citations": results}

class ExportRequest(BaseModel):
    messages: List[Dict[str, Any]]
    format: str = "pdf"
    title: str = "Legal Research Session"

@app.post("/export/session")
async def export_session(request: ExportRequest):
    from fastapi.responses import FileResponse
    import os
    import datetime
    
    os.makedirs("temp_uploads", exist_ok=True)
    stamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
    
    if request.format == "pdf":
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        filepath = f"temp_uploads/session_export_{stamp}.pdf"
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Add basic custom style
        speaker_style = ParagraphStyle(
            'Speaker',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            spaceAfter=6,
            textColor='blue'
        )
        
        Story = []
        Story.append(Paragraph(f"Lumina Copilot - {request.title}", styles['Title']))
        Story.append(Spacer(1, 12))
        
        for msg in request.messages:
            role = "User" if msg['role'] == 'user' else "Research Assistant"
            Story.append(Paragraph(role, speaker_style))
            # simple text processing
            content = str(msg['content']).replace('\n', '<br/>')
            Story.append(Paragraph(content, styles['Normal']))
            Story.append(Spacer(1, 12))
            
        doc.build(Story)
        return FileResponse(filepath, filename=f"{request.title}.pdf", media_type="application/pdf")
        
    elif request.format == "docx":
        from docx import Document
        filepath = f"temp_uploads/session_export_{stamp}.docx"
        doc = Document()
        doc.add_heading(f"Lumina Copilot - {request.title}", 0)
        
        for msg in request.messages:
            role = "User" if msg['role'] == 'user' else "Research Assistant"
            doc.add_heading(role, level=2)
            doc.add_paragraph(str(msg['content']))
            
        doc.save(filepath)
        return FileResponse(filepath, filename=f"{request.title}.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'.")

class ArgumentBuilderRequest(BaseModel):
    issue: str
    side: str = "petitioner"
    relevant_case_ids: Optional[List[str]] = []

@app.post("/argument-builder")
async def argument_builder(request: ArgumentBuilderRequest):
    try:
        from langchain_groq import ChatGroq
        from langchain_core.prompts import ChatPromptTemplate
        from langchain_core.output_parsers import JsonOutputParser
        from pydantic import BaseModel
        
        class ArgumentOutput(BaseModel):
            facts: str
            issues: str
            arguments: str
            prayer: str
            
        parser = JsonOutputParser(pydantic_object=ArgumentOutput)
        llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0.2)
        
        prompt = ChatPromptTemplate.from_template(
            "You are an expert Indian Commercial Court lawyer representing the {side}.\n"
            "Given the legal issue below, build a structured modular argument.\n"
            "Issue: {issue}\n\n{format_instructions}"
        )
        
        chain = prompt | llm | parser
        res = chain.invoke({
            "side": request.side,
            "issue": request.issue,
            "format_instructions": parser.get_format_instructions()
        })
        
        return res
    except Exception as e:
        logger.error(f"Argument Builder failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to build argument")

class DeadlineTrackerRequest(BaseModel):
    filing_date: str
    service_date: Optional[str] = None
    case_type: str
    basis: Optional[str] = None
    suit_type: Optional[str] = None

@app.post("/deadline-tracker")
async def deadline_tracker(request: DeadlineTrackerRequest):
    from datetime import datetime, timedelta
    import math

    basis = (request.basis or "filing_date").strip()
    if basis not in {"filing_date", "service_date"}:
        raise HTTPException(status_code=400, detail=f"Unsupported basis: {basis}")

    if basis == "service_date":
        if not request.service_date:
            raise HTTPException(status_code=400, detail="service_date is required when basis=service_date")
        date_to_parse = request.service_date
    else:
        date_to_parse = request.filing_date

    try:
        base_date = datetime.strptime(date_to_parse, "%Y-%m-%d")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid date format: {str(e)}")

    case_type = (request.case_type or "").strip()

    if case_type == "Original Suit":
        if basis != "service_date":
            raise HTTPException(
                status_code=400,
                detail="For Original Suit, statutory written statement timelines are triggered from service of summons. Select basis=service_date and provide service_date."
            )
        suit_type = (request.suit_type or "civil").strip().lower()
        if suit_type not in {"civil", "commercial"}:
            raise HTTPException(status_code=400, detail=f"Unsupported suit_type: {suit_type}")

        max_days = 120 if suit_type == "commercial" else 90
        rules = [
            {
                "name": "Written Statement (standard)",
                "days_offset": 30,
                "source": "Order VIII Rule 1 CPC (trigger: service of summons)",
            },
            {
                "name": "Written Statement (max)",
                "days_offset": max_days,
                "source": "Commercial Courts Act, 2015 read with Order VIII Rule 1 CPC (Commercial)" if suit_type == "commercial" else "CPC (extendable up to 90 days; subject to court discretion)",
            },
        ]
    else:
        raise HTTPException(
            status_code=501,
            detail="Only Original Suit written statement deadlines are implemented with statutory trigger logic. Add rule sets for other case types before enabling."
        )

    now = datetime.now()
    computed = []
    for rule in rules:
        due_date = base_date + timedelta(days=int(rule["days_offset"]))
        diff_seconds = (due_date - now).total_seconds()
        days_remaining = int(math.ceil(diff_seconds / 86400))
        expired = days_remaining < 0
        computed.append(
            {
                "name": rule["name"],
                "due_date": due_date.strftime("%Y-%m-%d"),
                "days_remaining": days_remaining,
                "urgency": "red" if days_remaining < 7 else ("amber" if days_remaining < 21 else "green"),
                "source": rule["source"],
                "basis": basis,
                "expired": expired,
            }
        )

    return {
        "deadlines": computed,
        "meta": {
            "basis": basis,
            "base_date": base_date.strftime("%Y-%m-%d"),
            "case_type": case_type,
            "suit_type": (request.suit_type or "civil").strip().lower() if case_type == "Original Suit" else None,
        },
    }


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
